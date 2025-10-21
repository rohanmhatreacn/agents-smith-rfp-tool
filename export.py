import logging
from datetime import datetime
from pathlib import Path
from typing import Any
from io import BytesIO
import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors

logger = logging.getLogger(__name__)


class Exporter:
    """
    Handles export of RFP proposal content to various formats.
    Supports DOCX, PDF, and image generation for diagrams.
    """
    
    def __init__(self):
        """Initializes exporter with default styling and formatting."""
        self.styles = getSampleStyleSheet()
    
    
    def export_docx(self, content: dict[str, Any], output_path: str) -> str:
        """
        Exports proposal content to a formatted DOCX document.
        Includes sections, tables, and embedded images.
        Returns the path to the created file.
        """
        try:
            doc = Document()
            
            doc.add_heading('RFP Proposal Response', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
            doc.add_paragraph()
            
            for section_name, section_data in content.items():
                if section_name in ['metadata', 'session_id']:
                    continue
                
                doc.add_heading(section_name.replace('_', ' ').title(), 1)
                
                if isinstance(section_data, dict):
                    self._add_dict_to_docx(doc, section_data)
                elif isinstance(section_data, str):
                    doc.add_paragraph(section_data)
                
                doc.add_paragraph()
            
            doc.save(output_path)
            logger.info(f"✅ Exported DOCX: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ Failed to export DOCX: {e}")
            raise
    
    
    def _add_dict_to_docx(self, doc: Document, data: dict):
        """
        Adds dictionary content to DOCX with proper formatting.
        Handles nested structures and tables.
        """
        for key, value in data.items():
            if isinstance(value, list) and value and isinstance(value[0], dict):
                self._add_table_to_docx(doc, value)
            elif isinstance(value, dict):
                doc.add_paragraph(f"{key}:", style='Heading 3')
                self._add_dict_to_docx(doc, value)
            else:
                doc.add_paragraph(f"{key}: {value}")
    
    
    def _add_table_to_docx(self, doc: Document, data: list[dict]):
        """
        Adds a table to DOCX from list of dictionaries.
        First dict keys become table headers.
        """
        if not data:
            return
        
        headers = list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header).title()
        
        for row_data in data:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                row_cells[i].text = str(row_data.get(header, ''))
    
    
    def export_pdf(self, content: dict[str, Any], output_path: str) -> str:
        """
        Exports proposal content to a formatted PDF document.
        Uses ReportLab for professional layout and styling.
        Returns the path to the created file.
        """
        try:
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            
            title = Paragraph("RFP Proposal Response", self.styles['Title'])
            story.append(title)
            story.append(Spacer(1, 0.2*inch))
            
            date_text = f"Generated: {datetime.now().strftime('%B %d, %Y')}"
            story.append(Paragraph(date_text, self.styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            
            for section_name, section_data in content.items():
                if section_name in ['metadata', 'session_id']:
                    continue
                
                heading = Paragraph(
                    section_name.replace('_', ' ').title(),
                    self.styles['Heading1']
                )
                story.append(heading)
                story.append(Spacer(1, 0.1*inch))
                
                if isinstance(section_data, dict):
                    self._add_dict_to_pdf(story, section_data)
                elif isinstance(section_data, str):
                    para = Paragraph(section_data, self.styles['Normal'])
                    story.append(para)
                
                story.append(Spacer(1, 0.2*inch))
            
            doc.build(story)
            logger.info(f"✅ Exported PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ Failed to export PDF: {e}")
            raise
    
    
    def _add_dict_to_pdf(self, story: list, data: dict):
        """
        Adds dictionary content to PDF with proper formatting.
        Handles nested structures and tables.
        """
        for key, value in data.items():
            if isinstance(value, list) and value and isinstance(value[0], dict):
                self._add_table_to_pdf(story, value)
            elif isinstance(value, dict):
                heading = Paragraph(f"{key}:", self.styles['Heading3'])
                story.append(heading)
                self._add_dict_to_pdf(story, value)
            else:
                text = f"<b>{key}:</b> {value}"
                para = Paragraph(text, self.styles['Normal'])
                story.append(para)
    
    
    def _add_table_to_pdf(self, story: list, data: list[dict]):
        """
        Adds a table to PDF from list of dictionaries.
        Formats with headers and alternating row colors.
        """
        if not data:
            return
        
        headers = list(data[0].keys())
        table_data = [[str(h).title() for h in headers]]
        
        for row in data:
            table_data.append([str(row.get(h, '')) for h in headers])
        
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 0.2*inch))
    
    
    def save_diagram_image(self, diagram_data: str | dict, output_path: str, format: str = "png") -> str:
        """
        Converts diagram code/JSON to static image format (PNG/JPG).
        Used for final export after diagram editing is complete.
        Returns the path to the created image file.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            if isinstance(diagram_data, dict):
                diagram_data = json.dumps(diagram_data, indent=2)
            
            img = Image.new('RGB', (800, 600), color='white')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 12)
            except:
                font = ImageFont.load_default()
            
            draw.text((20, 20), "Architecture Diagram", fill='black', font=font)
            draw.text((20, 50), str(diagram_data)[:500], fill='black', font=font)
            
            img.save(output_path, format.upper())
            logger.info(f"✅ Saved diagram image: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ Failed to save diagram image: {e}")
            raise


exporter = Exporter()

